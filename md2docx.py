"""
Markdown to Word (DOCX) Converter

A professional tool for converting Markdown documents to Word format with advanced
formatting capabilities including custom styles, headers, footers, images, tables,
and footnotes.

Dependencies:
    - python-docx: Word document manipulation
    - pandoc: Document conversion (external dependency)
"""

from typing import Optional, Dict, List, Tuple
from enum import Enum
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
import subprocess
import shutil
import re


class PaperSize(Enum):
    """Supported paper sizes for document layout."""
    LETTER = "letter"  # 8.5 x 11 inches (215.9 x 279.4 mm)
    LEGAL = "legal"    # 8.5 x 14 inches (215.9 x 355.6 mm)
    A4 = "a4"          # 8.27 x 11.69 inches (210 x 297 mm)


class DocumentStyle(Enum):
    """Predefined document style templates."""
    REPORT = "report"
    NOTE = "note"
    LETTER = "letter"
    MEMO = "memo"


class DocumentConfig:
    """
    Configuration for document conversion and styling.

    Attributes:
        style: Document style template
        paper_size: Paper size for the document
        author: Document author name
        date: Document date
        heading_colors: RGB color tuples for heading levels 1-3
        footer_text: Footer text for odd and even pages
        font_name: Base font family
        base_font_size: Base font size in points
        margins: Page margins (top, right, bottom, left) in cm
        line_spacing: Line spacing multiplier
        generate_toc: Whether to generate table of contents
        language: Document language code (e.g., 'en-US', 'fr-CA')
    """

    def __init__(
        self,
        style: DocumentStyle = DocumentStyle.REPORT,
        paper_size: PaperSize = PaperSize.LETTER,
        author: str = "",
        date: str = "",
        heading_colors: Dict[int, Tuple[int, int, int]] = None,
        footer_text: Dict[str, str] = None,
        font_name: str = "Arial",
        base_font_size: int = 12,
        margins: Tuple[float, float, float, float] = (2, 2, 2, 2),
        line_spacing: float = 1.0,
        generate_toc: bool = True,
        language: str = "en-US"
    ):
        self.style = style
        self.paper_size = paper_size
        self.author = author
        self.date = date
        self.heading_colors = heading_colors or {
            1: (37, 150, 190),
            2: (37, 150, 190),
            3: (37, 150, 190)
        }
        self.footer_text = footer_text or {
            "odd": "Page",
            "even": "Page"
        }
        self.font_name = font_name
        self.base_font_size = base_font_size
        self.margins = margins
        self.line_spacing = line_spacing
        self.generate_toc = generate_toc
        self.language = language

    @classmethod
    def create_report_style(cls, **kwargs):
        """Create a preconfigured style for reports."""
        default_config = {
            'style': DocumentStyle.REPORT,
            'paper_size': PaperSize.LETTER,
            'heading_colors': {
                1: (37, 150, 190),
                2: (37, 150, 190),
                3: (37, 150, 190)
            },
            'footer_text': {
                "odd": "Right text | Page",
                "even": "Page | Left text"
            }
        }
        default_config.update(kwargs)
        return cls(**default_config)

    @classmethod
    def create_note_style(cls, **kwargs):
        """Create a preconfigured style for internal notes."""
        default_config = {
            'style': DocumentStyle.NOTE,
            'paper_size': PaperSize.LEGAL,
            'heading_colors': {
                1: (70, 70, 70),
                2: (100, 100, 100),
                3: (130, 130, 130)
            },
            'footer_text': {
                "odd": "Internal Note | Page",
                "even": "Page | Internal Note"
            },
            'margins': (1.5, 1.5, 1.5, 1.5)
        }
        default_config.update(kwargs)
        return cls(**default_config)


class MarkdownToDocxConverter:
    """
    Convert Markdown documents to Word (DOCX) format with advanced formatting.

    This converter uses Pandoc for initial conversion and applies extensive
    post-processing to ensure professional formatting including:
    - Custom heading styles and colors
    - Image embedding with size adjustment
    - Table formatting with borders
    - Footnote styling
    - Custom headers and footers
    - Language settings
    """

    def __init__(self, config: DocumentConfig = None):
        """
        Initialize the converter.

        Args:
            config: Document configuration. If None, uses default settings.
        """
        self.config = config or DocumentConfig()
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Verify that Pandoc is installed on the system."""
        if not shutil.which('pandoc'):
            raise RuntimeError(
                "Pandoc is not installed. "
                "Please install it: https://pandoc.org/installing.html"
            )

    def _extract_title_from_markdown(self, content: str) -> str:
        """
        Extract the main title (H1) from markdown content.

        Args:
            content: Markdown document content

        Returns:
            The first H1 heading found, or "Untitled Document" if none exists
        """
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                return line[2:].strip()
        return "Untitled Document"

    def _extract_image_references(self, content: str) -> List[dict]:
        """
        Extract image references from markdown content.

        Args:
            content: Markdown document content

        Returns:
            List of dictionaries containing image metadata
        """
        image_pattern = r'!\[(.*?)\]\((.*?)\)'
        matches = re.findall(image_pattern, content)

        image_refs = []
        for alt_text, path in matches:
            path = path.strip()
            if path.startswith('img/'):
                path = path[4:]

            image_refs.append({
                'alt_text': alt_text,
                'path': path,
                'original_markdown': f'![{alt_text}]({path})'
            })

        print(f"Found {len(image_refs)} image references")
        return image_refs

    def _remove_image_references(self, content: str) -> str:
        """
        Replace image references with placeholders.

        Args:
            content: Markdown document content

        Returns:
            Content with images replaced by placeholders
        """
        image_pattern = r'!\[(.*?)\]\((.*?)\)'
        return re.sub(image_pattern, r'[IMAGE_PLACEHOLDER]', content)

    def _run_pandoc_conversion(self, input_path: Path, output_path: Path, title: str) -> None:
        """
        Run pandoc conversion with custom heading level mapping.

        This method uses a Lua filter to adjust heading levels so that
        Markdown H2 becomes Word Heading 1, H3 becomes Heading 2, etc.

        Args:
            input_path: Path to input markdown file
            output_path: Path to output docx file
            title: Document title
        """
        try:
            lua_path = self._create_lua_script(input_path)
            cmd = self._build_pandoc_command(input_path, output_path, title, lua_path)

            try:
                result = subprocess.run(cmd, check=True, capture_output=True, text=True)
                print("Pandoc conversion completed successfully")
            except subprocess.CalledProcessError as e:
                print(f"Pandoc conversion failed: {e.stderr}")
                raise
            finally:
                self._cleanup_lua_script(lua_path)

        except Exception as e:
            print(f"Error during pandoc conversion: {str(e)}")
            raise

    def _create_lua_script(self, input_path: Path) -> Path:
        """
        Create a temporary Lua script for pandoc filtering.

        The script adjusts heading levels to match the desired hierarchy.

        Args:
            input_path: Path to input file (used to determine script location)

        Returns:
            Path to the created Lua script
        """
        lua_script = """
        function Header(el)
            if el.level > 1 then
                el.level = el.level - 1
            end
            return el
        end
        """
        lua_path = input_path.parent / "adjust_headers.lua"
        with open(lua_path, 'w', encoding='utf-8') as f:
            f.write(lua_script)
        return lua_path

    def _build_pandoc_command(self, input_path: Path, output_path: Path,
                             title: str, lua_path: Path) -> List[str]:
        """
        Build the pandoc command with all necessary options.

        Args:
            input_path: Input markdown file
            output_path: Output docx file
            title: Document title
            lua_path: Path to Lua filter script

        Returns:
            Command as list of strings
        """
        cmd = [
            'pandoc',
            str(input_path),
            '-o', str(output_path),
            '-f', 'markdown',
            '-t', 'docx',
            '--wrap=none',
            '--columns=999',
            '--lua-filter=' + str(lua_path),
            '-M', f'title={title}',
            '-M', f'author={self.config.author}',
            '-M', f'date={self.config.date}'
        ]

        if self.config.generate_toc:
            cmd.extend(['--toc', '--number-sections'])

        return cmd

    def _cleanup_lua_script(self, lua_path: Path) -> None:
        """
        Remove temporary Lua script.

        Args:
            lua_path: Path to Lua script to delete
        """
        try:
            if lua_path.exists():
                lua_path.unlink()
        except Exception as e:
            print(f"Warning: Could not delete temporary Lua file: {e}")

    def convert(self,
                input_file: str,
                output_file: str,
                working_dir: Optional[str] = None,
                extra_args: Optional[List[str]] = None) -> None:
        """
        Convert a Markdown file to Word format.

        Args:
            input_file: Name of the input markdown file
            output_file: Name for the output docx file
            working_dir: Working directory (defaults to current directory)
            extra_args: Additional pandoc arguments (currently unused)
        """
        work_dir, input_path, output_path = self._setup_paths(
            input_file, output_file, working_dir
        )

        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")

        img_dir = self._create_image_directory(work_dir)
        content = self._read_markdown_content(input_path)
        document_title = self._extract_title_from_markdown(content)
        image_refs = self._extract_image_references(content)
        temp_md = self._create_temp_markdown(content, work_dir, input_path)

        try:
            self._run_pandoc_conversion(temp_md, output_path, document_title)
            self._post_process_document(output_path, document_title)
            self._insert_images(output_path, image_refs, work_dir)
            print(f"Conversion successful! File saved: {output_path}")
        finally:
            self._cleanup_temp_markdown(temp_md)

    def _setup_paths(self, input_file: str, output_file: str,
                     working_dir: Optional[str]) -> Tuple[Path, Path, Path]:
        """
        Setup and validate file paths.

        Args:
            input_file: Input filename
            output_file: Output filename
            working_dir: Working directory

        Returns:
            Tuple of (working_dir, input_path, output_path)
        """
        work_dir = Path(working_dir).resolve() if working_dir else Path.cwd()
        input_path = work_dir / input_file
        output_path = work_dir / output_file
        return work_dir, input_path, output_path

    def _create_image_directory(self, work_dir: Path) -> Path:
        """
        Create image directory if it doesn't exist.

        Args:
            work_dir: Working directory

        Returns:
            Path to image directory
        """
        img_dir = work_dir / "img"
        if not img_dir.exists():
            img_dir.mkdir(parents=True)
            print(f"Created image directory: {img_dir}")
        return img_dir

    def _read_markdown_content(self, input_path: Path) -> str:
        """
        Read markdown file content.

        Args:
            input_path: Path to markdown file

        Returns:
            File content as string
        """
        with open(input_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _create_temp_markdown(self, content: str, work_dir: Path,
                             input_path: Path) -> Path:
        """
        Create temporary markdown file without image references.

        Args:
            content: Original markdown content
            work_dir: Working directory
            input_path: Original input path

        Returns:
            Path to temporary file
        """
        temp_content = self._remove_image_references(content)
        temp_md = work_dir / f"temp_{input_path.name}"
        with open(temp_md, 'w', encoding='utf-8') as f:
            f.write(temp_content)
        return temp_md

    def _cleanup_temp_markdown(self, temp_md: Path) -> None:
        """
        Remove temporary markdown file.

        Args:
            temp_md: Path to temporary file
        """
        if temp_md.exists():
            temp_md.unlink()

    def _setup_footers(self, doc: Document) -> None:
        """
        Configure document footers with custom text and page numbers.

        Creates different footers for odd and even pages with page numbering.

        Args:
            doc: Document object to modify
        """
        for section in doc.sections:
            # Configure section for different odd/even pages
            sectPr = section._sectPr
            if not sectPr.find(qn('w:titlePg')):
                titlePg = OxmlElement('w:titlePg')
                sectPr.append(titlePg)
            if not sectPr.find(qn('w:evenAndOddHeaders')):
                evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
                sectPr.append(evenAndOddHeaders)

            # Configure headers/footers
            section.different_first_page_header_footer = True
            section.odd_and_even_pages_header_footer = True

            # Configure odd page footer (right pages)
            footer_odd = section.footer
            p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
            p_odd.clear()
            p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_odd.paragraph_format.space_before = Pt(12)

            # Add text and page number for odd pages
            run_odd = p_odd.add_run(f"{self.config.footer_text['odd']} | ")
            run_odd.font.name = self.config.font_name
            run_odd.font.size = Pt(10)
            self._add_page_number(p_odd)

            # Configure even page footer (left pages)
            footer_even = section.even_page_footer
            p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
            p_even.clear()
            p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_even.paragraph_format.space_before = Pt(12)

            # Add page number and text for even pages
            self._add_page_number(p_even)
            run_even = p_even.add_run(f" | {self.config.footer_text['even']}")
            run_even.font.name = self.config.font_name
            run_even.font.size = Pt(10)

    def _add_page_number(self, paragraph) -> None:
        """
        Add a dynamic page number field to a paragraph.

        Args:
            paragraph: Paragraph object to add page number to
        """
        run = paragraph.add_run()

        # Begin field character
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # Instruction text
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        # End field character
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

    def _process_footnotes(self, doc: Document) -> None:
        """
        Configure footnote formatting and language settings.

        Args:
            doc: Document object to modify
        """
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        try:
            # Configure footnote text style
            style = doc.styles['Footnote Text']
            style.font.name = self.config.font_name
            style.font.size = Pt(10)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0

            # Configure footnote reference style
            ref_style = doc.styles['Footnote Reference']
            ref_style.font.name = self.config.font_name
            ref_style.font.size = Pt(10)
            ref_style.font.superscript = True

            # Apply language to footnote style
            style_element = style._element
            rPr = style_element.get_or_add_rPr()

            # Remove existing language settings
            lang_elements = rPr.findall(qn('w:lang'))
            for lang in lang_elements:
                rPr.remove(lang)

            # Add new language setting
            lang = OxmlElement('w:lang')
            lang.set(qn('w:val'), self.config.language)
            lang.set(qn('w:eastAsia'), self.config.language)
            lang.set(qn('w:bidi'), self.config.language)
            rPr.append(lang)

            # Process each footnote if they exist
            if hasattr(doc, '_part') and hasattr(doc._part, '_footnotes_part') and doc._part._footnotes_part:
                footnotes = doc._part._footnotes_part.element.xpath('//w:footnote')
                for footnote in footnotes:
                    # Process each paragraph in the footnote
                    for p in footnote.xpath('.//w:p'):
                        pPr = p.get_or_add_pPr()

                        # Set paragraph properties
                        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>')
                        existing_spacing = pPr.find(qn('w:spacing'))
                        if existing_spacing is not None:
                            pPr.remove(existing_spacing)
                        pPr.append(spacing)

                        # Process runs in the footnote
                        for r in p.xpath('.//w:r'):
                            rPr = r.get_or_add_rPr()

                            # Remove all existing language settings
                            lang_elements = rPr.findall(qn('w:lang'))
                            for lang in lang_elements:
                                rPr.remove(lang)

                            # Add new language setting
                            lang = OxmlElement('w:lang')
                            lang.set(qn('w:val'), self.config.language)
                            lang.set(qn('w:eastAsia'), self.config.language)
                            lang.set(qn('w:bidi'), self.config.language)
                            rPr.append(lang)

        except Exception as e:
            print(f"Warning while processing footnotes: {str(e)}")

    def _post_process_document(self, doc_path: Path, document_title: str) -> None:
        """
        Post-process the Word document with all formatting requirements.

        This method applies all styling, removes duplicate titles, and
        processes all document elements.

        Args:
            doc_path: Path to the document file
            document_title: Main document title
        """
        doc = Document(doc_path)

        # Apply basic styles
        self._apply_global_styles(doc)

        # Find and remove any duplicate main title
        first_title_found = False
        paragraphs_to_remove = []

        for i, para in enumerate(doc.paragraphs):
            if para.text == document_title:
                if not first_title_found:
                    # Keep the first occurrence and style it as Title
                    first_title_found = True
                    para.style = doc.styles['Title']
                    para.clear()
                    run = para.add_run(document_title)
                    run.font.name = self.config.font_name
                    run.font.size = Pt(24)
                    run.font.bold = True
                else:
                    # Mark additional occurrences for removal
                    paragraphs_to_remove.append(i)

        # Remove duplicate titles (in reverse order to maintain indices)
        for idx in reversed(paragraphs_to_remove):
            p = doc.paragraphs[idx]._element
            p.getparent().remove(p)

        # Process all remaining paragraphs
        self._process_paragraphs(doc)
        self._process_tables(doc)
        self._process_footnotes(doc)
        self._setup_footers(doc)

        # Save the changes
        doc.save(doc_path)

    def _insert_images(self, doc_path: Path, image_refs: List[dict], work_dir: Path) -> None:
        """
        Insert images into the document at placeholder locations.

        Args:
            doc_path: Path to the document
            image_refs: List of image reference dictionaries
            work_dir: Working directory containing img folder
        """
        doc = Document(doc_path)
        img_dir = work_dir / "img"

        for para in doc.paragraphs:
            if '[IMAGE_PLACEHOLDER]' in para.text and image_refs:
                img_ref = image_refs.pop(0)
                image_path = img_dir / img_ref['path']

                print(f"Processing image: {image_path}")

                if image_path.exists():
                    para.clear()
                    run = para.add_run()
                    try:
                        picture = run.add_picture(str(image_path))

                        # Set image size with aspect ratio preservation
                        max_width = Inches(6)
                        if picture.width > max_width:
                            aspect_ratio = picture.height / picture.width
                            picture.width = max_width
                            picture.height = int(max_width * aspect_ratio)

                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"Added image: {img_ref['path']}")

                    except Exception as e:
                        print(f"Error adding image {img_ref['path']}: {str(e)}")
                        para.text = f"[Image: {img_ref['alt_text']}]"
                else:
                    print(f"Image not found: {image_path}")
                    para.text = f"[Image not found: {img_ref['alt_text']}]"

        doc.save(doc_path)

    def _process_paragraphs(self, doc: Document) -> None:
        """
        Process all paragraphs in the document with formatting and language.

        Args:
            doc: Document object to modify
        """
        for para in doc.paragraphs:
            # Skip paragraphs with specific styles
            if para.style.name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
                continue

            # Process paragraphs with 'Normal' style or no specific style
            if not para.style or para.style.name == 'Normal':
                # Apply base formatting to the paragraph
                para.style = doc.styles['Normal']
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.line_spacing = self.config.line_spacing

                # Process each run in the paragraph
                for run in para.runs:
                    # Apply base font settings
                    run.font.name = self.config.font_name
                    run.font.size = Pt(self.config.base_font_size)

                    # Ensure font is set even if not specified
                    if not run.font.name:
                        run.font.name = self.config.font_name
                    if not run.font.size:
                        run.font.size = Pt(self.config.base_font_size)

                    # Set language for this run
                    run_element = run._element
                    rPr = run_element.get_or_add_rPr()

                    # Remove all existing language settings
                    lang_elements = rPr.findall(qn('w:lang'))
                    for lang in lang_elements:
                        rPr.remove(lang)

                    # Add new language setting
                    lang = OxmlElement('w:lang')
                    lang.set(qn('w:val'), self.config.language)
                    lang.set(qn('w:eastAsia'), self.config.language)
                    lang.set(qn('w:bidi'), self.config.language)
                    rPr.append(lang)

    def _apply_global_styles(self, doc: Document) -> None:
        """
        Apply global document styles including language and formatting.

        Args:
            doc: Document object to modify
        """
        # Set document language
        element = OxmlElement('w:lang')
        element.set(qn('w:val'), self.config.language)
        element.set(qn('w:eastAsia'), self.config.language)
        element.set(qn('w:bidi'), self.config.language)

        styles_element = doc.styles.element
        if styles_element.find(qn('w:docDefaults')) is None:
            doc_defaults = OxmlElement('w:docDefaults')
            styles_element.insert(0, doc_defaults)
        else:
            doc_defaults = styles_element.find(qn('w:docDefaults'))

        if doc_defaults.find(qn('w:rPrDefault')) is None:
            r_pr_default = OxmlElement('w:rPrDefault')
            doc_defaults.insert(0, r_pr_default)
        else:
            r_pr_default = doc_defaults.find(qn('w:rPrDefault'))

        if r_pr_default.find(qn('w:rPr')) is None:
            r_pr = OxmlElement('w:rPr')
            r_pr_default.insert(0, r_pr)
        else:
            r_pr = r_pr_default.find(qn('w:rPr'))

        # Remove existing language
        existing_lang = r_pr.find(qn('w:lang'))
        if existing_lang is not None:
            r_pr.remove(existing_lang)

        # Add new language
        r_pr.append(element)

        # Configure standard styles
        styles_to_configure = {
            'Normal': {
                'font_size': self.config.base_font_size,
                'font_name': self.config.font_name,
                'bold': False,
                'space_before': 0,
                'space_after': 0,
                'line_spacing': self.config.line_spacing
            },
            'Title': {
                'font_size': 24,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 12,
                'space_after': 12,
                'line_spacing': 1.0
            }
        }

        # Configure heading styles (## = Heading 1, ### = Heading 2, #### = Heading 3)
        heading_styles = {
            'Heading 1': {
                'font_size': 18,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 18,
                'space_after': 12,
                'line_spacing': 1.0,
                'color': self.config.heading_colors.get(1, (0, 0, 0)),
                'italic': False
            },
            'Heading 2': {
                'font_size': 16,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 16,
                'space_after': 10,
                'line_spacing': 1.0,
                'color': self.config.heading_colors.get(2, (0, 0, 0)),
                'italic': False
            },
            'Heading 3': {
                'font_size': 14,
                'font_name': self.config.font_name,
                'bold': True,
                'space_before': 14,
                'space_after': 8,
                'line_spacing': 1.0,
                'color': self.config.heading_colors.get(3, (0, 0, 0)),
                'italic': True
            }
        }

        # Merge all styles configurations
        styles_to_configure.update(heading_styles)

        # Apply configurations to each style
        for style_name, config in styles_to_configure.items():
            try:
                style = doc.styles[style_name]

                # Font configuration
                style.font.name = config['font_name']
                style.font.size = Pt(config['font_size'])
                style.font.bold = config['bold']
                if config.get('italic'):
                    style.font.italic = True

                # Color (if specified)
                if 'color' in config:
                    style.font.color.rgb = RGBColor(*config['color'])

                # Paragraph formatting
                style.paragraph_format.space_before = Pt(config['space_before'])
                style.paragraph_format.space_after = Pt(config['space_after'])
                style.paragraph_format.line_spacing = config['line_spacing']

                # For heading styles, ensure they're not linked to other styles
                if style_name.startswith('Heading'):
                    if hasattr(style, 'base_style'):
                        style.base_style = None

                    # Remove any existing numbering
                    if hasattr(style._element, 'pPr'):
                        pPr = style._element.pPr
                        if pPr is not None:
                            numPr = pPr.find(qn('w:numPr'))
                            if numPr is not None:
                                pPr.remove(numPr)

            except KeyError:
                print(f"Style '{style_name}' not found")

        # Configure section properties
        for section in doc.sections:
            # Set page size
            if self.config.paper_size == PaperSize.LETTER:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
            elif self.config.paper_size == PaperSize.LEGAL:
                section.page_width = Inches(8.5)
                section.page_height = Inches(14)
            else:  # A4
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)

            # Set margins
            section.top_margin = Cm(self.config.margins[0])
            section.right_margin = Cm(self.config.margins[1])
            section.bottom_margin = Cm(self.config.margins[2])
            section.left_margin = Cm(self.config.margins[3])

    def _process_tables(self, doc: Document) -> None:
        """
        Process all tables in the document with formatting and borders.

        Args:
            doc: Document object to modify
        """
        table_font_size = 10  # Fixed size for table content

        for table in doc.tables:
            # Process header row (first row)
            if table.rows:
                for cell in table.rows[0].cells:
                    for para in cell.paragraphs:
                        # Configure header cell
                        for run in para.runs:
                            run.font.name = self.config.font_name
                            run.font.size = Pt(table_font_size)
                            run.font.bold = True

                        # Ensure paragraph has at least one run
                        if not para.runs:
                            run = para.add_run()
                            run.font.name = self.config.font_name
                            run.font.size = Pt(table_font_size)
                            run.font.bold = True

            # Process all rows
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Configure regular cell
                        for run in para.runs:
                            run.font.name = self.config.font_name
                            run.font.size = Pt(table_font_size)

                            # Set language for this run
                            run_element = run._element
                            rPr = run_element.get_or_add_rPr()

                            # Remove all existing language settings
                            lang_elements = rPr.findall(qn('w:lang'))
                            for lang in lang_elements:
                                rPr.remove(lang)

                            # Add new language setting
                            lang = OxmlElement('w:lang')
                            lang.set(qn('w:val'), self.config.language)
                            lang.set(qn('w:eastAsia'), self.config.language)
                            lang.set(qn('w:bidi'), self.config.language)
                            rPr.append(lang)

                        # Ensure paragraph has proper formatting
                        if not para.runs:
                            run = para.add_run()
                            run.font.name = self.config.font_name
                            run.font.size = Pt(table_font_size)

                        # Set paragraph properties
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)
                        para.paragraph_format.line_spacing = 1.0

                    # Add borders to cell
                    self._set_cell_borders(cell)

    def _set_cell_borders(self, cell: _Cell) -> None:
        """
        Add borders to a table cell.

        Args:
            cell: Cell object to add borders to
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Create borders element
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

        # Add each border
        for border in ['top', 'left', 'bottom', 'right']:
            edge = OxmlElement(f'w:{border}')
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), '4')
            edge.set(qn('w:space'), '0')
            edge.set(qn('w:color'), 'auto')
            tcBorders.append(edge)


def main():
    """Example usage demonstrating the converter capabilities."""
    # Example configuration for a professional report
    config = DocumentConfig.create_report_style(
        author="Your Name",
        date="2025-01-31",
        language="en-US",
        generate_toc=True,
        paper_size=PaperSize.LETTER,
        font_name="Arial",
        base_font_size=12,
        heading_colors={
            1: (37, 150, 190),
            2: (37, 150, 190),
            3: (37, 150, 190)
        },
        footer_text={
            "odd": "Document Title",
            "even": "Author Name"
        }
    )

    # Create converter and process file
    converter = MarkdownToDocxConverter(config)

    try:
        converter.convert(
            input_file="example.md",
            output_file="example.docx",
            working_dir="."
        )
    except Exception as e:
        print(f"Error during conversion: {str(e)}")


if __name__ == "__main__":
    main()
